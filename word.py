from docx import Document
from ast import literal_eval


async def create_survey_docx(surveys):

    doc = Document()

    for survey in surveys:
        doc.add_heading('                                                        SO‘ROVNOMA', level=1)
        # Parse subjects and weekdays safely
        subjects = literal_eval(survey.subjects)
        weekdays = literal_eval(survey.weekdays)

        # Add paragraph with student info
        doc.add_paragraph(
            f'Men {survey.first_name} {survey.last_name} “{survey.classroom[:-1]}-{survey.classroom[-1]}” sinf o’quvchisi 2024–2025 o’quv yili tugagach, '
            f'dolzarb 90 kunlik ta’tilda ota-onam roziligi bilan quyidagi to‘garaklarga qatnashmoqchiman:'
        )

        # Add the table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        if survey.subjects and survey.weekdays and survey.classroom and survey.date:

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Fanlar'
            hdr_cells[1].text = 'Qayerda joylashgan'
            hdr_cells[2].text = 'O’quv markaz nomi'
            hdr_cells[3].text = 'Qachondan bormoqdasiz'
            hdr_cells[4].text = 'Haftaning qaysi kunlari'

            # Data row
            row_cells = table.add_row().cells
            row_cells[0].text = ', '.join(subjects)
            row_cells[1].text = survey.place
            row_cells[2].text = survey.education_name
            row_cells[3].text = str(survey.date)
            row_cells[4].text = ', '.join(weekdays)

            # Add a space after each user
            doc.add_paragraph()

        # Save once for all users
    file_path = "14-school.docx"
    doc.save(file_path)
    return file_path
